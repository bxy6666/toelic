from __future__ import annotations

import textwrap
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "toeic-dev-process-screenshots.docx"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/consola.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


FONT_TITLE = font(32, True)
FONT_SUBTITLE = font(22, True)
FONT_BODY = font(20)
FONT_MONO = font(18)
FONT_SMALL = font(16)


def read_text(path: str) -> str:
    return (ROOT / path).read_text(encoding="utf-8", errors="replace")


def wrap_line(line: str, width: int) -> list[str]:
    if not line:
        return [""]
    chunks: list[str] = []
    current = ""
    for ch in line:
        ch_width = 2 if ord(ch) > 127 else 1
        cur_width = sum(2 if ord(c) > 127 else 1 for c in current)
        if cur_width + ch_width > width:
            chunks.append(current)
            current = ch
        else:
            current += ch
    chunks.append(current)
    return chunks


def make_text_screenshot(
    filename: str,
    title: str,
    subtitle: str,
    lines: list[str],
    *,
    width: int = 1500,
    max_lines: int = 30,
    accent: tuple[int, int, int] = (18, 126, 102),
) -> Path:
    wrapped: list[str] = []
    for line in lines:
        wrapped.extend(wrap_line(line.rstrip(), 92))
    wrapped = wrapped[:max_lines]

    line_h = 30
    height = 190 + max(1, len(wrapped)) * line_h + 54
    img = Image.new("RGB", (width, height), (248, 250, 252))
    draw = ImageDraw.Draw(img)

    draw.rounded_rectangle((28, 24, width - 28, height - 24), radius=22, fill=(255, 255, 255), outline=(219, 234, 231), width=2)
    draw.rounded_rectangle((28, 24, width - 28, 118), radius=22, fill=(235, 252, 247), outline=(219, 234, 231), width=2)
    draw.rectangle((28, 92, width - 28, 118), fill=(235, 252, 247))
    draw.rectangle((28, 118, width - 28, 124), fill=accent)

    draw.text((58, 46), title, fill=(15, 23, 42), font=FONT_TITLE)
    draw.text((58, 86), subtitle, fill=(51, 65, 85), font=FONT_SMALL)

    y = 150
    for line in wrapped:
        color = (15, 23, 42)
        if line.startswith("#") or line.startswith("##"):
            color = accent
            f = FONT_SUBTITLE
        elif line.startswith("|") or line.startswith("- ") or line.startswith("`") or "PASS" in line or "DONE" in line:
            f = FONT_MONO
        else:
            f = FONT_BODY
        draw.text((58, y), line, fill=color, font=f)
        y += line_h

    path = ASSET_DIR / filename
    img.save(path)
    return path


def select_lines(text: str, starts: list[str], count: int) -> list[str]:
    lines = text.splitlines()
    selected: list[str] = []
    for marker in starts:
        for idx, line in enumerate(lines):
            if marker in line:
                selected.extend(lines[idx : idx + count])
                selected.append("")
                break
    return selected


def build_assets() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for wrong in ["codex-current-conversation.png", "codex-conversation-window.png"]:
        p = ASSET_DIR / wrong
        if p.exists():
            p.unlink()

    spec = read_text("docs/spec-driven/spec.md")
    design = read_text("docs/spec-driven/design.md")
    tasks = read_text("docs/spec-driven/tasks.md")
    acceptance = read_text("docs/spec-driven/acceptance.md")
    changelog = read_text("docs/change-log-ai.md")

    assets: dict[str, Path] = {}
    assets["dialogue_note"] = make_text_screenshot(
        "01-model-dialogue-note.png",
        "大模型对话截图说明",
        "以下对话图来自 Codex 本地会话日志摘录；如课堂要求 UI 原图，可再从聊天界面补拍",
        [
            "已补入：5 月 8 日规格生成、5 月 9 日 MaaS 测试、公网修复、图片 API 方案、5 月 10 日汇报材料任务的会话摘录。",
            "",
            "如需要补拍聊天 UI 原图，建议补充：",
            "1. 用户确认 spec.md 进入 Gate 2 的对话。",
            "2. 用户确认 design.md 进入 Gate 3 的对话。",
            "3. 用户确认 tasks.md 进入 Gate 4 并开始 T01-T22 的对话。",
            "4. MaaS 真实生成、UI 迭代、公网验收等关键反馈对话。",
            "",
            "注意：补拍时避开真实 API Key、Authorization、Bearer Token 和 .env.local 内容。",
        ],
        accent=(217, 119, 6),
    )
    dialogue_1 = [
        "2026-05-08 用户：我准备采用 spec_driven coding 的方式来开发一个网页，主要功能包括听力练习、语法练习、错题本、学习统计、题目由大模型生成。",
        "2026-05-08 Codex：先按规则读取全局/项目规则和现有缓存，然后只生成 spec.md，不启动开发、不碰业务代码。",
        "2026-05-08 Codex：项目级规则和缓存都不存在，触发轻量初始化；只看根目录结构和工具可用性，不安装依赖、不启动脚手架。",
        "2026-05-08 Codex：准备写 spec.md，并补极简 docs/project-context.md 和 docs/change-log-ai.md，记录当前为空白规划项目。",
        "2026-05-08 Codex：spec.md 已生成，未启动开发流程，等待用户审核。",
    ]
    assets["dialogue_1"] = make_text_screenshot(
        "01a-model-dialogue-spec.png",
        "大模型对话截图 1",
        "从用户需求到 spec.md：先规划、再审核、不直接开发",
        dialogue_1,
        accent=(217, 119, 6),
    )
    dialogue_2 = [
        "2026-05-09 用户：目前已完成对大模型 apikey 的配置，帮我测试各个题目是否能正确生成，并给出下一步完善计划。",
        "2026-05-09 Codex：只针对题目生成链路做最小范围测试，不顺手改业务代码。",
        "2026-05-09 Codex：项目缓存可复用，MaaS 生成链路已有历史通过记录；本次不完整重扫，只找生成接口、题型入口和测试方式。",
        "2026-05-09 Codex：定位核心接口 app/api/ai/generate-questions/route.ts，会按 listening / grammar 调 MaaS、校验结构并落库。",
        "2026-05-09 Codex：确认配置文件存在 API Key、Base URL 和模型名；后续汇报不回显密钥。",
    ]
    assets["dialogue_2"] = make_text_screenshot(
        "01b-model-dialogue-maas.png",
        "大模型对话截图 2",
        "MaaS 真实生成测试：定位接口、校验字段、保护密钥",
        dialogue_2,
        accent=(217, 119, 6),
    )
    dialogue_3 = [
        "2026-05-09 用户：公网创建的网站未能将本地运行的所有页面正确上传连接，仅仅显示背景板没显示完整功能。",
        "2026-05-09 Codex：只定位“本地正常、公网页面只剩背景”的发布差异点，先缩小到路由、静态资源路径、构建产物或部署入口配置。",
        "2026-05-09 Codex：项目缓存显示这是 Next.js + Prisma/SQLite 混合项目，最近加了 Cloudflare Quick Tunnel。",
        "2026-05-09 Codex：API 调用都是相对路径 /api/...，对公网隧道是正确的；接下来排除 node_modules，只看项目文件并用浏览器复现。",
        "2026-05-09 Codex：确认首页、错题页、统计页、设置页都会服务端访问 SQLite/Prisma，因此外部平台部署不会自动拥有本机数据库和 Key。",
    ]
    assets["dialogue_3"] = make_text_screenshot(
        "01c-model-dialogue-public.png",
        "大模型对话截图 3",
        "公网问题排查：从现象到路由、隧道、数据库边界判断",
        dialogue_3,
        accent=(217, 119, 6),
    )
    dialogue_4 = [
        "2026-05-09 用户：听力练习图片描述风格无法生成具体图片如何解决，调用新 API 生成图像吗。",
        "2026-05-09 Codex：当前链路是 DeepSeek-V3.2 文本生成，图片描述题本质仍是文字题；需要核对官方图像 API 口径。",
        "2026-05-09 Codex：推荐接新图像 API，DeepSeek 继续生成题目 JSON，OpenAI Image API 只负责生成原创场景图。",
        "2026-05-09 用户：PLEASE IMPLEMENT THIS PLAN: 听力图片描述真实配图方案。",
        "2026-05-09 结果：后续变更记录显示已新增 imageUrl/imagePrompt 字段、图像生成服务、Prompt 强化、迁移、lint/build 验证。",
    ]
    assets["dialogue_4"] = make_text_screenshot(
        "01d-model-dialogue-image-api.png",
        "大模型对话截图 4",
        "图片 API 方案：先判断边界，再接入最小真实配图链路",
        dialogue_4,
        accent=(217, 119, 6),
    )
    assets["spec"] = make_text_screenshot(
        "02-spec-driven-gates.png",
        "需求规格与 Gate 流程截图",
        "spec.md - 先规格后实现，状态为 Approved",
        select_lines(spec, ["状态：Approved", "## 0. Spec-Driven 流程", "## 1. 项目概述"], 12),
    )
    assets["design"] = make_text_screenshot(
        "03-design-architecture.png",
        "设计文档截图",
        "design.md - 架构、页面、数据库、API、MaaS 与测试策略",
        select_lines(design, ["# TOEIC Practice Studio Design", "## 1.", "MaaS", "Prisma"], 10),
    )
    task_lines = select_lines(tasks, ["## 1. 任务总览"], 28)
    assets["tasks"] = make_text_screenshot(
        "04-task-breakdown-t01-t22.png",
        "任务拆分截图",
        "tasks.md - T01 到 T22 分步推进，状态均为 DONE",
        task_lines,
        max_lines=34,
    )
    assets["acceptance"] = make_text_screenshot(
        "05-acceptance-record.png",
        "验收记录截图",
        "acceptance.md - lint / build / smoke / generation / public acceptance",
        select_lines(acceptance, ["# TOEIC Practice Studio Acceptance", "MaaS", "npm run", "PASS"], 12),
        accent=(37, 99, 235),
    )
    log_lines = []
    for line in changelog.splitlines():
        if line.startswith("## 2026-05-08 Gate") or line.startswith("## 2026-05-08 V") or line.startswith("## 2026-05-09") or line.startswith("## 2026-05-10"):
            log_lines.append(line)
        if len(log_lines) >= 22:
            break
    assets["changelog"] = make_text_screenshot(
        "06-ai-change-log-timeline.png",
        "AI 变更日志截图",
        "docs/change-log-ai.md - 每轮实现、修复与验证都有记录",
        log_lines,
        max_lines=28,
        accent=(124, 58, 237),
    )

    tree_lines = [
        "E:\\托业",
        "|-- app/                 Next.js 页面与 API Routes",
        "|-- components/          AppShell、练习区、统计面板、UI 组件",
        "|-- lib/                 MaaS Client、题目校验、统计、错题、设置服务",
        "|-- prisma/              SQLite 数据模型与 migration",
        "|-- prompts/             听力 / 语法题目生成 Prompt",
        "|-- scripts/smoke/       浏览器、generation、公网验收脚本",
        "|-- scripts/public/      公网启动与 tunnel 脚本",
        "|-- scripts/reports/     课程汇报与过程文档生成脚本",
        "|-- docs/                文档索引、上下文与 AI 变更记录",
        "|-- docs/spec-driven/    需求规格、设计、任务拆分、验收记录",
    ]
    assets["tree"] = make_text_screenshot(
        "07-project-structure.png",
        "工程结构截图",
        "目录结构证明代码、文档、脚本与验收材料分层清晰",
        tree_lines,
        accent=(15, 118, 110),
    )
    validation_lines = [
        "PS E:\\托业> npm run lint",
        "PASS - ESLint 检查通过",
        "",
        "PS E:\\托业> npm run build",
        "PASS - Next.js 生产构建通过",
        "",
        "PS E:\\托业> npm run smoke:generation",
        "PASS - 9 个听力 / 语法子题型真实生成通过",
        "",
        "PS E:\\托业> node scripts/smoke/browser-smoke.mjs",
        "PASS - 核心页面浏览器 smoke 通过",
        "",
        "PS E:\\托业> npm run smoke:public",
        "PASS - 公网地址 41 项验收通过",
    ]
    assets["validation"] = make_text_screenshot(
        "08-terminal-validation-summary.png",
        "终端验证截图",
        "由 acceptance.md 与 change-log-ai.md 汇总的关键验证命令",
        validation_lines,
        accent=(22, 101, 52),
    )
    return assets


def add_image(doc: Document, image_path: Path, caption: str, width: float = 6.25) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = "Caption"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"


def build_docx(assets: dict[str, Path]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    styles["Caption"].font.name = "Microsoft YaHei"
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.color.rgb = RGBColor(71, 85, 105)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TOEIC Practice Studio 开发过程截图说明")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(22)

    intro = doc.add_paragraph()
    intro.add_run("范围：").bold = True
    intro.add_run("仅收录开发过程截图与说明，不包含课堂展示 PPT 页面内容。材料覆盖 spec-driven 流程、任务拆分、AI 变更记录、工程验证、产品运行和响应式验收。")

    add_heading(doc, "1. 大模型对话截图", 1)
    add_image(doc, assets["dialogue_note"], "说明：对话图来自 Codex 本地会话日志摘录；如课堂要求聊天窗口原图，可按本页提示补拍。")
    add_image(doc, assets["dialogue_1"], "截图说明：用户先要求采用 spec-driven coding，Codex 明确只生成规格，不提前开发。")
    add_image(doc, assets["dialogue_2"], "截图说明：API Key 配好后，Codex 只测试 MaaS 题目生成链路，并避免回显密钥。")
    add_image(doc, assets["dialogue_3"], "截图说明：公网问题排查围绕路由、隧道、Prisma/SQLite 和运行环境边界展开。")
    add_image(doc, assets["dialogue_4"], "截图说明：图片描述题从文本题边界分析，推进到最小图像 API 接入方案。")

    add_heading(doc, "2. Spec-Driven 过程", 1)
    add_image(doc, assets["spec"], "截图说明：需求规格已进入 Approved 状态，并明确 spec -> design -> tasks -> implementation -> acceptance 的 Gate 流程。")
    add_image(doc, assets["design"], "截图说明：设计阶段先确定架构、数据库、API、MaaS 调用、错误处理和测试策略，再进入实现。")
    add_image(doc, assets["tasks"], "截图说明：T01-T22 将开发拆成可验收小任务，体现逐项推进和范围控制。")
    add_image(doc, assets["acceptance"], "截图说明：验收记录用于闭环 lint、build、真实生成、浏览器和公网验证结果。")
    add_image(doc, assets["changelog"], "截图说明：AI 变更记录保留每轮实现、修复、验证和交付材料生成过程。")

    add_heading(doc, "3. 工程实现与验证", 1)
    add_image(doc, assets["tree"], "截图说明：工程目录将页面、组件、服务、Prisma、Prompt、脚本和文档分层管理。")
    add_image(doc, assets["validation"], "截图说明：关键验证命令均有记录，包括 lint、build、题型生成、浏览器 smoke 和公网验收。")

    add_heading(doc, "4. 产品运行过程截图", 1)
    product_images = [
        ("output/playwright/v1-2-home.png", "V1.2 首页动效与视觉重设计后的运行截图。"),
        ("output/playwright/v1-3-listening.png", "V1.3 听力练习流程完善，包含生成、答题和结果状态。"),
        ("output/playwright/v1-4-home-desktop.png", "V1.4 桌面端视觉对齐与响应式巡检。"),
        ("output/playwright/v1-4-home-mobile.png", "V1.4 移动端首页布局检查。"),
        ("output/playwright/v1-6-home-desktop.png", "V1.6 首页 Hero 今日计划卡片迭代结果。"),
        ("output/playwright/v1-7-stats-desktop.png", "V1.7 学习统计卡通仪表盘桌面端截图。"),
        ("output/playwright/v1-7-stats-mobile.png", "V1.7 学习统计卡通仪表盘移动端截图。"),
        ("output/playwright/public-desktop-listening.png", "公网验收：听力练习页桌面端截图。"),
        ("output/playwright/public-desktop-grammar.png", "公网验收：语法练习页桌面端截图。"),
        ("output/playwright/public-desktop-mistakes.png", "公网验收：错题本页面桌面端截图。"),
        ("output/playwright/public-desktop-settings.png", "公网验收：设置页仅展示配置状态，不暴露完整 API Key。"),
        ("output/playwright/public-mobile-listening.png", "公网验收：听力练习页移动端截图。"),
        ("output/playwright/public-mobile-stats.png", "公网验收：统计页移动端截图。"),
    ]
    for rel, caption in product_images:
        path = ROOT / rel
        if path.exists():
            add_image(doc, path, caption)

    add_heading(doc, "5. 交付说明", 1)
    notes = [
        "本文档未放入 output/presentation 下的 PPT 预览页或 PPT 内容截图。",
        "第 1 节的大模型对话截图来自 Codex 本地会话日志摘录，不包含真实密钥；如课堂要求聊天 UI 原图，可按说明补拍替换。",
        "建议补拍对话时避开真实 API Key、Authorization、Bearer Token、.env.local 等敏感内容。",
    ]
    for item in notes:
        doc.add_paragraph(item, style="List Bullet")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def main() -> None:
    assets = build_assets()
    build_docx(assets)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
